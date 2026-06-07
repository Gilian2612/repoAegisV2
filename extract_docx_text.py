from docx import Document
import os

def read_docx(file_path):
    doc = Document(file_path)
    parts = []

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text.strip())

    for table in doc.tables:
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    cells.append(cell_text)
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)

if __name__ == "__main__":
    intake_path = os.path.join("input", "intake.docx")
    estate_path = os.path.join("input", "estate_plan.docx")

    if not os.path.exists(intake_path):
        print(f"ERROR: Missing file: {intake_path}")
        print("Make sure the intake file is inside the input folder and named intake.docx")
        exit()

    if not os.path.exists(estate_path):
        print(f"ERROR: Missing file: {estate_path}")
        print("Make sure the estate plan file is inside the input folder and named estate_plan.docx")
        exit()

    intake_text = read_docx(intake_path)
    estate_text = read_docx(estate_path)

    os.makedirs("output", exist_ok=True)

    with open("output/intake_text.txt", "w", encoding="utf-8") as f:
        f.write(intake_text)

    with open("output/estate_plan_text.txt", "w", encoding="utf-8") as f:
        f.write(estate_text)

    print("Text extracted successfully.")
    print("Created: output/intake_text.txt")
    print("Created: output/estate_plan_text.txt")